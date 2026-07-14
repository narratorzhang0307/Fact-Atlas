#!/usr/bin/env python3
"""Build the final bilingual FactRelay form-copy and handoff document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_video_script import (
    BLUE,
    CYAN,
    INK,
    LIME,
    MUTED,
    PAPER,
    PINK,
    VIOLET,
    VIOLET_SOFT,
    YELLOW,
    configure_document,
    set_font,
    set_table_geometry,
    shade,
)


ROOT = Path(__file__).resolve().parents[1]
DELIVERY = ROOT / "FactRelay_黑客松交付包"
OUT = DELIVERY / "01_提交材料" / "FactRelay_报名表填写稿_中英双语.docx"

BRIEF = "https://hackathonweekly.feishu.cn/wiki/M0pewmd0ti3z8IkVmVYcEzTWnIe"
FORM = "https://hackathonweekly.feishu.cn/share/base/form/shrcnaF4yF8HmLhA42AH7DWlrPc"
DEMO = "https://factrelay-ai3-2026.yediqizhang37.chatgpt.site"
REPO = "https://github.com/narratorzhang0307/FactRelay"
VIDEO_RELEASE = "https://github.com/narratorzhang0307/FactRelay/releases/tag/ai3-2026-submission"
VIDEO_DIRECT = "https://github.com/narratorzhang0307/FactRelay/releases/download/ai3-2026-submission/FactRelay_Demo_2m30s_Bilingual.mp4"

PITCH_CN = "FactRelay 是一个基于 Gonka 的可追溯多模型事实核查工作台，将公开证据、对抗审查、确定性 Truth Score 与真实推理回执合并成一份可审计结论。"
PITCH_EN = "FactRelay is a traceable multi-model fact-checking workbench on Gonka that turns public evidence, adversarial review, a deterministic Truth Score, and real inference receipts into one auditable verdict."

INTRO_CN = (
    "社交媒体与生成式 AI 让信息生成变得廉价，却让“为什么应该相信这个结论”变得更难回答。大多数 AI 事实核查产品仍只返回一段自信的文字，用户看不到来源是否真正回应主张，也看不到不同模型之间的分歧。\n\n"
    "FactRelay 将文本、公开链接或截图转化为一次可追溯调查。系统先检索当前公开证据，再通过 GonkaRouter 让 Kimi-K2.6 担任调查方，让 MiniMax-M2.7 以质疑方角色检查循环引用、时间错位、因果跳跃和遗漏背景。最终 Truth Score 由确定性代码计算，并与来源账本、双模型判断、分歧程度、执行路径和原始 Gonka Request ID 一起展示。\n\n"
    "FactRelay 不把推理回执当成“真理证明”；回执只证明哪一次 Gonka 请求生成了分析，事实仍由可检查的公开证据支持。"
)

INTRO_EN = (
    "Social media and generative AI have made information cheap to produce, while making one question harder to answer: why should anyone trust the verdict? Most AI fact checkers still return one confident paragraph, hiding whether the evidence addresses the exact claim and whether independent models disagreed.\n\n"
    "FactRelay turns text, a public URL, or a screenshot into a traceable investigation. It retrieves current public evidence, sends the evidence packet to Kimi-K2.6 as the investigator, and asks MiniMax-M2.7 to act as an adversarial skeptic looking for source laundering, chronology errors, causal leaps, and omitted context. Tested deterministic code computes the Truth Score and keeps the evidence ledger, model disagreement, execution path, and untouched upstream Gonka Request IDs visible.\n\n"
    "A request receipt proves which Gonka call produced the analysis; factual support still comes from inspectable evidence."
)

TECH_CN = (
    "React 前端通过 Node/Cloudflare Worker API 接入核查链路。输入首先经过类型、长度、图片大小和 URL/SSRF 防护；非 AI 检索层抓取公开页面，并并发请求 Google News RSS 与 Bing News RSS。内置长城示例还透明地实时抓取 NASA、ESA、Smithsonian 权威页面，以抵抗新闻搜索服务偶发不可用。\n\n"
    "所有语义推理均通过 https://api.gonkarouter.io/v1/chat/completions 完成。Kimi 负责主张提取和调查，MiniMax 负责对抗审查。模型只能引用已检索账本中的来源编号，越界编号在评分前被代码拒绝。Truth Score = 50 + 50 ×（55% 模型共识 + 45% 来源加权证据）；弱证据会把评分拉回 50，模型分歧会降低结论信心。若模型偶发返回非法 JSON，系统仅对同一 Gonka 模型进行一次严格结构化重试，并将失败请求保留为 partial 轨迹。"
)

TECH_EN = (
    "The React client calls a Node/Cloudflare Worker verification API. Input passes type, length, image-size, URL, redirect, DNS, and SSRF guards. A non-AI retrieval layer fetches public pages and races Google News RSS with Bing News RSS. The built-in Great Wall starter also fetches a transparent live allowlist of NASA, ESA, and Smithsonian pages.\n\n"
    "All semantic inference uses the GonkaRouter chat-completions endpoint. Kimi extracts/investigates; MiniMax performs adversarial review. Models may cite only retrieved source indexes. Deterministic code computes Truth Score = 50 + 50 × (55% model consensus + 45% source-weighted evidence), pulls weak evidence toward uncertainty, and lowers confidence when models disagree. One same-model structured retry handles malformed JSON while preserving the failed call as a partial trace step."
)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_heading_pair(doc: Document, english: str, chinese: str, level: int = 1) -> None:
    heading = doc.add_heading(level=level)
    r = heading.add_run(english)
    set_font(r, None, True, BLUE if level == 1 else INK)
    r2 = heading.add_run(f"  /  {chinese}")
    set_font(r2, None, True, VIOLET)


def add_body(doc: Document, text: str, color: str = INK, size: float = 10.2) -> None:
    for block in text.split("\n\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.line_spacing = 1.24
        r = p.add_run(block)
        set_font(r, size, False, color)


def add_link_box(doc: Document, label: str, url: str, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "\n")
    set_font(r, 7.5, True, VIOLET)
    r = p.add_run(url)
    set_font(r, 9.2, True, INK)


def add_field_table(doc: Document, rows: list[tuple[str, str]], widths=(2280, 7080)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_geometry(table, list(widths))
    headers = table.rows[0].cells
    shade(headers[0], VIOLET)
    shade(headers[1], VIOLET)
    for cell, text in zip(headers, ("FORM FIELD · 表单字段", "PASTE THIS · 可直接粘贴")):
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, 8, True, "FFFFFF")
    set_repeat_table_header(table.rows[0])
    fills = [VIOLET_SOFT, PAPER, LIME, PAPER, CYAN, PAPER, YELLOW, PAPER, PINK]
    for index, (label, value) in enumerate(rows):
        cells = table.add_row().cells
        shade(cells[0], fills[index % len(fills)])
        shade(cells[1], "FFFFFF")
        p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(label)
        set_font(r0, 8.5, True, INK)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.18
        r1 = p1.add_run(value)
        set_font(r1, 9, False, INK)


def add_checklist(doc: Document, items: list[str], checked: bool = True) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.05)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(("☑" if checked else "☐") + "  " + item)
        set_font(r, 9.6, False, INK)


def configure_submission_document(doc: Document) -> None:
    configure_document(doc)
    section = doc.sections[0]
    hp = section.header.paragraphs[0]
    hp.clear()
    r = hp.add_run("FACTRELAY  /  FINAL SUBMISSION PACK  ·  最终报名材料")
    set_font(r, 7.5, True, MUTED)
    doc.core_properties.title = "FactRelay 黑客松报名表填写稿"
    doc.core_properties.subject = "AI³ Growth Hackathon 2026 · Track 3 · Gonka"
    doc.core_properties.author = "FactRelay"


def build(output: Path) -> None:
    doc = Document()
    configure_submission_document(doc)

    p = doc.add_paragraph()
    r = p.add_run("AI³ GROWTH HACKATHON 2026  ·  TRACK 3  ·  GONKA: AI FOR SOCIETY")
    set_font(r, 8.5, True, VIOLET)
    title = doc.add_paragraph(style="Title")
    r = title.add_run("FactRelay")
    set_font(r, 32, True, INK)
    p = doc.add_paragraph()
    r = p.add_run("Final Form Copy & Delivery Guide")
    set_font(r, 20, True, INK)
    p = doc.add_paragraph()
    r = p.add_run("报名表填写稿与最终交付说明 · 中英双语")
    set_font(r, 13, True, MUTED)

    cover = DELIVERY / "02_演示视频" / "FactRelay_Demo_Cover.png"
    if cover.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(cover), width=Inches(6.45))

    add_link_box(doc, "Official submission form · 官方提交表", FORM, LIME)
    add_link_box(doc, "Running demo · 在线 Demo", DEMO, CYAN)
    add_link_box(doc, "Public repository · 公开仓库", REPO, VIOLET_SOFT)
    add_link_box(doc, "2:30 demo video · 2分30秒视频", VIDEO_RELEASE, PINK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("Status / 状态：")
    set_font(r, 9, True, VIOLET)
    r = p.add_run("所有公开材料已完成；仅姓名、手机号、邮箱和最终授权确认需参赛者本人填写。")
    set_font(r, 9, False, INK)

    add_page_break(doc)
    add_heading_pair(doc, "Official form map", "官方表单字段")
    add_body(doc, "The form has ten required fields. The public fields below are complete; private contact and consent fields stay blank by design.\n\n官方表单共有 10 个必填项。公开字段已完成；私人联系信息和授权确认按设计留空。", MUTED)
    add_field_table(doc, [
        ("1 · 项目名", "FactRelay"),
        ("2 · 队伍成员姓名", "【参赛者填写】队长写第一个"),
        ("3 · 手机号", "【参赛者填写】确保组委会可以联系"),
        ("4 · 邮箱", "【参赛者填写】"),
        ("5 · 赛道", "Gonka: AI for Society"),
        ("6 · Github 链接", REPO),
        ("7 · 作品体验网址/名称", f"FactRelay — {DEMO}"),
        ("8 · 作品介绍/技术说明", "见本文件第 3–5 页，可直接复制中文稿；如表单字数有限，优先复制“一句话介绍 + 中文项目介绍 + 中文技术说明”。"),
        ("9 · 产品演示视频链接", VIDEO_RELEASE + "\n直链：" + VIDEO_DIRECT),
        ("10 · 授权确认", "【参赛者本人阅读并选择】好，我已知晓并同意 / 否，不参赛"),
    ])

    add_heading_pair(doc, "Official requirements", "官方材料要求", level=2)
    add_checklist(doc, [
        "可运行 Demo：公开网址已上线，Gonka liveReady=true。",
        "公开 GitHub：README 包含功能、安装、运行和接入方式。",
        "产品演示视频：150 秒，中文旁白，烧录英文字幕。",
        "技术说明：赛道、Sponsor 技术、架构、关键功能和迭代计划齐全。",
        "Gonka 专项：在线 Demo 展示两个模型的真实 Gonka Request ID。",
    ])

    add_page_break(doc)
    add_heading_pair(doc, "One-line pitch", "一句话介绍")
    add_heading_pair(doc, "Chinese", "中文", level=2)
    add_body(doc, PITCH_CN, size=11)
    add_heading_pair(doc, "English", "英文", level=2)
    add_body(doc, PITCH_EN, MUTED, 10.5)

    add_heading_pair(doc, "Project introduction", "项目介绍")
    add_heading_pair(doc, "Chinese — recommended form copy", "中文推荐粘贴稿", level=2)
    add_body(doc, INTRO_CN)
    add_heading_pair(doc, "English", "英文", level=2)
    add_body(doc, INTRO_EN, MUTED)

    add_heading_pair(doc, "Technical description", "技术说明")
    add_heading_pair(doc, "Chinese — recommended form copy", "中文推荐粘贴稿", level=2)
    add_body(doc, TECH_CN)
    add_heading_pair(doc, "English", "英文", level=2)
    add_body(doc, TECH_EN, MUTED)

    add_heading_pair(doc, "Architecture at a glance", "架构速览", level=2)
    architecture = [
        ("INPUT", "Text / URL / image · 输入防护与 SSRF 校验", LIME),
        ("RETRIEVAL", "Public HTML + Google/Bing News RSS · 非 AI 实时证据", YELLOW),
        ("KIMI", "Investigator + image claim extraction · Gonka 调查方", CYAN),
        ("MINIMAX", "Adversarial skeptic · Gonka 质疑方", VIOLET_SOFT),
        ("OUTPUT", "Truth Score + evidence ledger + request receipts · 可追溯输出", PINK),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_geometry(table, [1872] * 5)
    for cell, (label, text, fill) in zip(table.rows[0].cells, architecture):
        shade(cell, fill)
        p = cell.paragraphs[0]
        r = p.add_run(label + "\n")
        set_font(r, 7.5, True, INK)
        r = p.add_run(text)
        set_font(r, 8, False, INK)

    add_heading_pair(doc, "Gonka compliance", "Gonka 赛道合规")
    add_checklist(doc, [
        "所有 AI 推理与信息语义验证逻辑均通过 gonkarouter.io 运行。",
        "模型一：moonshotai/Kimi-K2.6，负责图片主张提取与调查。",
        "模型二：MiniMaxAI/MiniMax-M2.7，负责对抗式交叉审查。",
        "输入支持文本、公开链接与图片；输出 0–100 Truth Score。",
        "界面展示完整推理轨迹与每一步真实 Gonka Request ID。",
        "检索层不调用其他 AI；失败时不会伪造结论或回执。",
    ])

    add_heading_pair(doc, "Live production proof", "公开站真实运行证明", level=2)
    add_field_table(doc, [
        ("Date / 日期", "2026-07-15"),
        ("Mode / 模式", "live · 实时案"),
        ("Verdict / 结论", "Refuted · 事实不符"),
        ("Truth Score", "18 / 100"),
        ("Decision confidence", "88%"),
        ("Sources / 来源", "5 retrievable public sources"),
        ("Gonka receipts", "Kimi 与 MiniMax 均返回非空上游 Request ID"),
    ], widths=(2700, 6660))

    screenshot = DELIVERY / "03_截图" / "08_真实结果_04_Gonka回执.png"
    if screenshot.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(screenshot), width=Inches(5.0))
        p = doc.add_paragraph("真实公开运行的回执卡片 / Live public receipt card")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_font(run, 8, False, MUTED)

    add_heading_pair(doc, "Verification report", "交付验收")
    add_field_table(doc, [
        ("Code quality", "TypeScript clean · 19/19 tests · production build passes"),
        ("Dependency audit", "npm audit --audit-level=low · 0 known vulnerabilities"),
        ("Public demo", "HTTP 200 · liveReady=true · browser console clean"),
        ("Video", "150.000s · 1920×1080 · 30fps · H.264 · AAC 48kHz stereo"),
        ("Video privacy", "No API key, email, phone, password, or unrelated tab visible"),
        ("Public access", "Release page and 6.0MB direct MP4 work without authentication"),
    ])

    add_heading_pair(doc, "Submission sequence", "最终提交步骤", level=2)
    add_checklist(doc, [
        f"打开官方提交表：{FORM}",
        "填写队伍成员姓名、手机号和邮箱。",
        "赛道选择 Gonka: AI for Society。",
        "复制本文件的 GitHub、Demo、介绍/技术说明和视频链接。",
        "本人阅读授权说明，并选择“好，我已知晓并同意”。",
        "提交后打开“查看提交记录”，确认链接和字段均完整。",
    ], checked=False)

    add_heading_pair(doc, "Private fields — participant only", "私人字段，仅参赛者填写", level=2)
    add_field_table(doc, [
        ("队伍成员姓名", "____________________________________________"),
        ("手机号", "____________________________________________"),
        ("邮箱", "____________________________________________"),
        ("授权确认", "☐ 已阅读并同意赛事方授权说明"),
    ])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("Official brief / 官方说明：")
    set_font(r, 8, True, VIOLET)
    r = p.add_run(BRIEF)
    set_font(r, 8, False, MUTED)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    build(OUT)
