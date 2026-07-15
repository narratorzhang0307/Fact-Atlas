#!/usr/bin/env python3
"""Build professional, plain-language FactRelay subtitle masters without touching the video."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_video_script import add_field, set_font, set_table_geometry, shade


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "FactRelay_黑客松交付包" / "02_演示视频"
DOCX_OUT = OUT_DIR / "FactRelay_正式字幕稿_中英双语.docx"
CN_SRT_OUT = OUT_DIR / "FactRelay_Chinese_Subtitles.srt"
EN_SRT_OUT = OUT_DIR / "FactRelay_English_Subtitles.srt"
BILINGUAL_SRT_OUT = OUT_DIR / "FactRelay_Bilingual_Subtitles.srt"

INK = "11120F"
MUTED = "5F645E"
VIOLET = "7865FF"
BLUE = "3C49DA"
LIME = "B8FF5C"
CYAN = "BCECF2"
YELLOW = "FFE38A"
PINK = "FFBED8"
PAPER = "F8F7F2"
VIOLET_SOFT = "EDE9FF"
WHITE = "FFFFFF"

SUBTITLES = [
    ("00:00:00,000", "00:00:05,800", "AI 事实核查，不该让我们再相信另一个黑盒。", "AI fact checking should not replace one black box with another."),
    ("00:00:05,800", "00:00:12,000", "FactRelay 把每次核查变成一条可翻看、可追溯的证据链。", "FactRelay turns each investigation into a navigable evidence chain."),
    ("00:00:12,000", "00:00:19,500", "可以提交文本、公开链接或社交媒体截图。", "Submit text, a public link, or a social-media screenshot."),
    ("00:00:19,500", "00:00:27,000", "实时核查会提取准确主张，并检索最新公开证据。", "A live run extracts the claim, then retrieves current public evidence."),
    ("00:00:27,000", "00:00:33,800", "这里核查一个常见说法：从月球上能肉眼看到长城。", "We test the familiar claim that the Great Wall is visible from the Moon."),
    ("00:00:33,800", "00:00:41,000", "FactRelay 先收集证据，再依次调用两位 Gonka 模型。", "FactRelay collects evidence before calling two Gonka models in sequence."),
    ("00:00:41,000", "00:00:48,500", "结论是“事实不符”，并给出可复算的 Truth Score。", "The verdict is Refuted, with a deterministic Truth Score."),
    ("00:00:48,500", "00:00:58,000", "分数由模型共识和来源加权证据共同计算。", "The score is deterministic: model consensus plus source-weighted evidence."),
    ("00:00:58,000", "00:01:07,000", "每条来源都保留发布者、日期、链接、立场和可信度。", "Every source keeps its publisher, date, URL, stance, and reliability."),
    ("00:01:07,000", "00:01:16,000", "模型只能引用这份编号清单；无效来源编号会被拒绝。", "Models may cite only this numbered ledger; invalid source indexes are rejected."),
    ("00:01:16,000", "00:01:25,500", "Kimi-K2.6 担任调查方，先形成基于证据的判断。", "Kimi-K2.6 investigates and forms the first evidence-based judgment."),
    ("00:01:25,500", "00:01:36,000", "MiniMax-M2.7 担任质疑方，检查循环引用、时间错误和遗漏背景。", "MiniMax-M2.7 challenges it for source laundering, chronology errors, and missing context."),
    ("00:01:36,000", "00:01:44,500", "每次推理都保留 Gonka 上游 Request ID 和执行顺序。", "Every inference preserves its upstream Gonka Request ID and execution order."),
    ("00:01:44,500", "00:01:54,000", "回执证明分析来自哪次请求，但不等于事实本身已被证明。", "A receipt proves which call produced the analysis—not whether the claim is true."),
    ("00:01:54,000", "00:02:02,000", "所有语义推理只通过 GonkaRouter 运行。", "All semantic inference runs exclusively through GonkaRouter."),
    ("00:02:02,000", "00:02:10,000", "可测试评分、SSRF 防护和空预览回执，让系统边界清楚可见。", "Tested scoring, SSRF guards, and null preview receipts make the boundary explicit."),
    ("00:02:10,000", "00:02:17,000", "结论、来源、审查和回执，每一部分都能单独检查。", "Verdict, Sources, Review, and Proof remain individually inspectable."),
    ("00:02:17,000", "00:02:24,000", "用户得到的不是一句答案，而是一条可以复核的证据链。", "The user receives a reviewable chain—not just an answer."),
    ("00:02:24,000", "00:02:30,000", "FactRelay：质疑主张，保留回执。", "FactRelay. Question the claim. Keep the receipts."),
]


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        normal._element.rPr.rFonts.set(qn(key), "Arial Unicode MS")

    for name, size, color, before, after in (
        ("Title", 26, INK, 0, 6),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Arial Unicode MS"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            style._element.rPr.rFonts.set(qn(key), "Arial Unicode MS")

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.add_run("FACTRELAY  /  SUBTITLE MASTER  ·  正式字幕母版"), 7.5, True, MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    set_font(fp.add_run("AI³ GROWTH HACKATHON 2026   ·   "), 7.5, True, MUTED)
    add_field(fp, "PAGE")

    doc.core_properties.title = "FactRelay 正式字幕稿（中英双语）"
    doc.core_properties.subject = "2分30秒演示视频独立字幕母版"
    doc.core_properties.author = "FactRelay"
    doc.core_properties.keywords = "FactRelay, Gonka, subtitles, SRT, bilingual"


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def write_srt(path: Path, mode: str) -> None:
    blocks = []
    for index, (start, end, chinese, english) in enumerate(SUBTITLES, start=1):
        if mode == "cn":
            text = chinese
        elif mode == "en":
            text = english
        else:
            text = f"{chinese}\n{english}"
        blocks.append(f"{index}\n{start} --> {end}\n{text}")
    path.write_text("\n\n".join(blocks) + "\n", encoding="utf-8")


def build_docx(path: Path) -> None:
    doc = Document()
    configure_document(doc)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_font(kicker.add_run("AI³ GROWTH HACKATHON 2026  ·  TRACK 3"), 8.5, True, VIOLET)

    title = doc.add_paragraph(style="Title")
    set_font(title.add_run("FactRelay 2:30 正式字幕母版"), 26, True, INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("Professional, plain-language bilingual subtitles / 专业、准确、通俗易懂"), 11.5, True, MUTED)

    metrics = doc.add_table(rows=1, cols=4)
    metrics.style = "Table Grid"
    metric_data = [
        ("02:30", "总时长", VIOLET_SOFT),
        ("19", "字幕条目", LIME),
        ("CN / EN", "双语母版", CYAN),
        ("SRT + DOCX", "交付格式", YELLOW),
    ]
    for cell, (value, label, fill) in zip(metrics.rows[0].cells, metric_data):
        shade(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(value + "\n"), 12, True, INK)
        set_font(p.add_run(label), 7.5, True, MUTED)
    set_table_geometry(metrics, [2340, 2340, 2340, 2340])

    doc.add_paragraph()
    note = doc.add_table(rows=1, cols=1)
    note.style = "Table Grid"
    cell = note.cell(0, 0)
    shade(cell, PAPER)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("使用说明 / REVIEW FIRST\n"), 8, True, VIOLET)
    set_font(p.add_run("这是独立字幕文件，尚未重新烧录视频。中文优先保证准确、自然和易懂；英文保持简洁，适合画面内两行显示。"), 10.5, False, INK)
    set_table_geometry(note, [9360])

    doc.add_heading("Time-coded subtitle master / 带时间码字幕总表", level=1)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    set_font(intro.add_run("时间码已与 150 秒最终视频对齐。中文用于审阅或中文字幕轨；英文可用于国际评委字幕轨。"), 10, False, MUTED)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for cell, label in zip(headers, ("TIME / 时间", "中文（通俗专业）", "ENGLISH")):
        shade(cell, VIOLET)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), 8.5, True, WHITE)
    set_repeat_header(table.rows[0])

    row_colors = [VIOLET_SOFT, LIME, CYAN, YELLOW, PINK]
    for index, (start, end, chinese, english) in enumerate(SUBTITLES, start=1):
        cells = table.add_row().cells
        shade(cells[0], row_colors[(index - 1) % len(row_colors)])
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(0)
        short_start = start[3:].replace(",", ".")
        short_end = end[3:].replace(",", ".")
        set_font(p0.add_run(f"{index:02d}\n{short_start}\n→ {short_end}"), 7.8, True, INK)
        for cell, value in ((cells[1], chinese), (cells[2], english)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.18
            set_font(p.add_run(value), 9.1, False, INK)
    set_table_geometry(table, [1320, 3900, 4140])

    doc.add_heading("Terminology / 术语说明", level=1)
    terms = [
        ("Truth Score", "保留英文产品名；中文可解释为“真实度评分”，强调它由代码计算。"),
        ("Request ID", "保留英文技术名；中文可称“请求回执”，只证明分析来自哪次调用。"),
        ("SSRF", "字幕保留缩写，口头可解释为“恶意链接访问防护”。"),
        ("GonkaRouter", "保持官方产品名，不翻译。"),
    ]
    for term, explanation in terms:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(term + "  "), 10, True, BLUE)
        set_font(p.add_run(explanation), 10, False, INK)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    write_srt(CN_SRT_OUT, "cn")
    write_srt(EN_SRT_OUT, "en")
    write_srt(BILINGUAL_SRT_OUT, "bilingual")
    build_docx(DOCX_OUT)
    for path in (DOCX_OUT, CN_SRT_OUT, EN_SRT_OUT, BILINGUAL_SRT_OUT):
        print(path)


if __name__ == "__main__":
    main()
